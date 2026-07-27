#!/usr/bin/env python3
"""Generate the Google Docs-ready CodeDecay hackathon submission artifact."""

from __future__ import annotations

import argparse
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_SOURCE = ROOT / "docs" / "hackathon" / "submission-description.md"
LEDGER_SOURCE = ROOT / "docs" / "hackathon" / "originality-ledger.md"
OUTPUT = ROOT / "docs" / "hackathon" / "codedecay-submission.docx"

INK = "000000"
MUTED = "555555"
BORDER = "DADCE0"
CODE_FILL = "F5F5F5"
CONTENT_WIDTH_DXA = 9360
TABLE_LABEL_DXA = 2160
TABLE_VALUE_DXA = CONTENT_WIDTH_DXA - TABLE_LABEL_DXA
INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|<https?://[^>]+>)"
)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate the CodeDecay hackathon submission DOCX."
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="Fail when the checked-in DOCX differs from a fresh deterministic build.",
    )
    args = parser.parse_args()

    if args.check:
        with tempfile.TemporaryDirectory(prefix="codedecay_submission_check_") as temp_dir:
            candidate = Path(temp_dir) / OUTPUT.name
            generate_document(candidate)
            if not OUTPUT.exists() or read_docx_parts(candidate) != read_docx_parts(OUTPUT):
                raise SystemExit(
                    "Checked-in hackathon DOCX is stale. "
                    "Run python3 scripts/generate-hackathon-submission-docx.py."
                )
        print(f"Verified {OUTPUT}")
        return

    generate_document(OUTPUT)
    print(f"Wrote {OUTPUT}")


def generate_document(output: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_markdown(doc, SUBMISSION_SOURCE.read_text(encoding="utf-8"))
    doc.add_page_break()
    add_originality_ledger(doc, LEDGER_SOURCE.read_text(encoding="utf-8"))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    normalize_docx_archive(output)


def read_docx_parts(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path, "r") as archive:
        return {name: archive.read(name) for name in sorted(archive.namelist())}


def normalize_docx_archive(path: Path) -> None:
    """Make ZIP metadata stable so CI can compare the generated DOCX byte-for-byte."""
    with zipfile.ZipFile(path, "r") as source:
        members = [
            (member.filename, source.read(member.filename), member.external_attr)
            for member in source.infolist()
        ]

    normalized = path.with_suffix(f"{path.suffix}.tmp")
    with zipfile.ZipFile(
        normalized,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        for filename, data, external_attr in sorted(members):
            member = zipfile.ZipInfo(filename, date_time=(1980, 1, 1, 0, 0, 0))
            member.compress_type = zipfile.ZIP_DEFLATED
            member.create_system = 3
            member.external_attr = external_attr
            target.writestr(member, data, compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)
    normalized.replace(path)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    core = doc.core_properties
    core.title = "CodeDecay — Codex India Hackathon 2026"
    core.subject = "Building Evals submission description and originality evidence"
    core.author = "SubmuxHQ"
    core.keywords = "CodeDecay, Codex, Building Evals, PR safety, hackathon"
    core.comments = (
        "Generated from docs/hackathon/submission-description.md and "
        "docs/hackathon/originality-ledger.md"
    )

    normal = doc.styles["Normal"]
    set_style_font(normal, "Arial", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    configure_heading(doc.styles["Heading 1"], 20, INK, 20, 6)
    configure_heading(doc.styles["Heading 2"], 16, INK, 18, 6)
    configure_heading(doc.styles["Heading 3"], 14, "434343", 16, 4)
    sanitize_unused_title_style(doc)

    for style_name in ("CodeDecay List Bullet", "CodeDecay List Number"):
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        set_style_font(style, "Arial", 11, INK)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.keep_together = True

    code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, "Courier New", 9.5, INK)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)

    caption = doc.styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(caption, "Arial", 9.5, MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.15
    (
        doc._codedecay_bullet_abstract_id,
        doc._codedecay_decimal_abstract_id,
    ) = configure_numbering(doc)


def sanitize_unused_title_style(doc: Document) -> None:
    """Avoid Google Docs importing a rule from Word's unused Title style."""
    title = doc.styles["Title"].element
    paragraph_properties = title.find(qn("w:pPr"))
    if paragraph_properties is not None:
        for border in list(paragraph_properties.findall(qn("w:pBdr"))):
            paragraph_properties.remove(border)
    run_properties = title.find(qn("w:rPr"))
    if run_properties is not None:
        for underline in list(run_properties.findall(qn("w:u"))):
            run_properties.remove(underline)


def configure_heading(style, size: float, color: str, before: float, after: float) -> None:
    set_style_font(style, "Arial", size, color)
    style.font.bold = False
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.keep_with_next = True


def set_style_font(style, family: str, size: float, color: str) -> None:
    style.font.name = family
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), family)
    fonts.set(qn("w:hAnsi"), family)
    fonts.set(qn("w:eastAsia"), family)


def add_markdown(doc: Document, markdown: str, *, skip_title: bool = False) -> None:
    lines = markdown.splitlines()
    paragraph_lines: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = " ".join(part.strip() for part in paragraph_lines).strip()
        paragraph_lines.clear()
        if text:
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, text)

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2)
            if skip_title and level == 1:
                index += 1
                continue
            if level == 1:
                add_title(doc, text)
            else:
                paragraph = doc.add_paragraph(style=f"Heading {level - 1}")
                add_inline_markdown(paragraph, text)
            index += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            items, index = collect_list_items(lines, index, "- ")
            num_id = add_numbering_instance(doc, doc._codedecay_bullet_abstract_id)
            for item in items:
                paragraph = doc.add_paragraph(style="CodeDecay List Bullet")
                apply_numbering(paragraph, num_id)
                add_inline_markdown(paragraph, item)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            items, index = collect_numbered_items(lines, index)
            num_id = add_numbering_instance(doc, doc._codedecay_decimal_abstract_id)
            for item in items:
                paragraph = doc.add_paragraph(style="CodeDecay List Number")
                apply_numbering(paragraph, num_id)
                add_inline_markdown(paragraph, item)
            continue

        paragraph_lines.append(stripped)
        index += 1

    flush_paragraph()


def collect_list_items(lines: list[str], index: int, marker: str) -> tuple[list[str], int]:
    items: list[str] = []
    current = ""
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            break
        if stripped.startswith(marker):
            if current:
                items.append(current)
            current = stripped[len(marker) :].strip()
        elif current and not stripped.startswith("#") and not stripped.startswith("```"):
            current = f"{current} {stripped}"
        else:
            break
        index += 1
    if current:
        items.append(current)
    return items, index


def collect_numbered_items(lines: list[str], index: int) -> tuple[list[str], int]:
    items: list[str] = []
    current = ""
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            break
        match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if match:
            if current:
                items.append(current)
            current = match.group(1)
        elif current and not stripped.startswith("#") and not stripped.startswith("```"):
            current = f"{current} {stripped}"
        else:
            break
        index += 1
    if current:
        items.append(current)
    return items, index


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, "Arial", 26, INK, bold=False)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for position, line in enumerate(lines):
        paragraph = doc.add_paragraph(style="Code Block")
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = position < len(lines) - 1
        if position == 0:
            paragraph.paragraph_format.space_before = Pt(4)
        if position == len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(8)
        shade_paragraph(paragraph, CODE_FILL)
        paragraph.add_run(line or " ")


def add_inline_markdown(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Courier New", 9.5, INK)
        elif token.startswith("["):
            label, url = re.match(r"^\[([^\]]+)\]\(([^)]+)\)$", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    run_properties.extend((fonts, color, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((run_properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run_font(
    run,
    family: str,
    size: float,
    color: str,
    *,
    bold: bool | None = None,
) -> None:
    run.font.name = family
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), family)
    fonts.set(qn("w:hAnsi"), family)
    fonts.set(qn("w:eastAsia"), family)


def configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    next_abstract = max(abstract_ids, default=0) + 1
    bullet = add_numbering_definition(
        numbering,
        abstract_id=next_abstract,
        number_format="bullet",
        marker="●",
    )
    decimal = add_numbering_definition(
        numbering,
        abstract_id=next_abstract + 1,
        number_format="decimal",
        marker="%1.",
    )
    return bullet, decimal


def add_numbering_definition(
    numbering,
    *,
    abstract_id: int,
    number_format: str,
    marker: str,
) -> int:
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), number_format)
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), marker)
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    paragraph_properties.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def add_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    number_properties = properties.find(qn("w:numPr"))
    if number_properties is None:
        number_properties = OxmlElement("w:numPr")
        properties.append(number_properties)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.extend((level, number))


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_originality_ledger(doc: Document, markdown: str) -> None:
    before_table, after_heading = markdown.split("## Hackathon-period ledger", 1)
    table_block, after_table = after_heading.split("## Evidence interpretation", 1)

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Appendix A: Originality and Codex Evidence Ledger")
    add_markdown(doc, before_table, skip_title=True)

    ledger_heading = doc.add_paragraph(style="Heading 2")
    ledger_heading.add_run("Hackathon-period ledger")
    rows = parse_markdown_table(table_block)
    add_ledger_table(doc, rows)

    post = "## Evidence interpretation" + after_table
    add_markdown(doc, post)


def parse_markdown_table(markdown: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        values = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        rows.append(values)
    if len(rows) < 2:
        raise ValueError("Originality ledger table is missing rows")
    return rows


def add_ledger_table(doc: Document, rows: list[list[str]]) -> None:
    headers, evidence_rows = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_geometry(table, [TABLE_LABEL_DXA, TABLE_VALUE_DXA])
    set_table_borders(table, BORDER)
    set_table_cell_margins(table, top=80, bottom=80, start=120, end=120)
    set_repeat_table_header(table.rows[0])
    set_cell_text(table.rows[0].cells[0], "Field", bold=True)
    set_cell_text(table.rows[0].cells[1], "Evidence", bold=True)

    for entry, values in enumerate(evidence_rows, start=1):
        separator = table.add_row()
        merged = separator.cells[0].merge(separator.cells[1])
        set_cell_text(merged, f"Entry {entry} — {strip_markdown(values[0])}", bold=True)
        keep_cell_with_next(merged)
        for label, value in zip(headers[1:], values[1:]):
            row = table.add_row()
            set_cell_text(row.cells[0], strip_markdown(label), bold=True)
            set_cell_markdown(row.cells[1], value)

    set_table_geometry(table, [TABLE_LABEL_DXA, TABLE_VALUE_DXA])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    citation = doc.add_paragraph(style="Table Citation")
    citation.add_run(
        "Source: public CodeDecay issues, pull requests, commits, CI checks, "
        "and deployment evidence linked in each entry."
    )


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, "Arial", 10, INK, bold=bold)


def set_cell_markdown(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    add_inline_markdown(paragraph, text)
    for run in paragraph.runs:
        if run.font.name != "Courier New":
            set_run_font(run, "Arial", 10, INK, bold=run.bold)


def strip_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text


def set_table_geometry(table, widths: list[int]) -> None:
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths)))

    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_cell_width(cell, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.append(cell_width)
    cell_width.set(qn("w:type"), "dxa")
    cell_width.set(qn("w:w"), str(width))


def set_table_cell_margins(
    table,
    *,
    top: int,
    bottom: int,
    start: int,
    end: int,
) -> None:
    properties = table._tbl.tblPr
    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.append(margins)
    for edge, width in (
        ("top", top),
        ("bottom", bottom),
        ("start", start),
        ("end", end),
    ):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def keep_cell_with_next(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.keep_with_next = True


if __name__ == "__main__":
    main()
